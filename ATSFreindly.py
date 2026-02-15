import re
import nltk
import spacy
from nltk.corpus import stopwords
from docx import Document
import pdfplumber
from sklearn.feature_extraction.text import CountVectorizer
import numpy as np

# Download the NLTK stopwords if not already installed
nltk.download('stopwords')
stop_words = set(stopwords.words('english'))

# Load spaCy's large English model for semantic analysis
nlp = spacy.load('en_core_web_sm')  # Use the smaller model

def read_docx(file_path):
    """ Read a .docx file and return the text content """
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def read_pdf(file_path):
    """ Read a .pdf file and return the text content using pdfplumber """
    with pdfplumber.open(file_path) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
        return text

def clean_text(text):
    """ Basic text cleaning (remove non-alphabetic, extra spaces) """
    text = re.sub(r'\s+', ' ', text)  # remove extra spaces
    text = re.sub(r'[^a-zA-Z\s]', '', text)  # remove non-alphabetical characters
    return text.lower()

def generate_ngrams(text, n=2):
    """ Generate n-grams from text """
    words = text.split()
    return [' '.join(words[i:i+n]) for i in range(len(words)-n+1)]

def check_keywords(resume_text, job_description_text):
    """ Check for the presence of job keywords in the resume """
    resume_words = set(clean_text(resume_text).split())
    job_desc_words = set(clean_text(job_description_text).split())
    
    # Removing stopwords from both sets to focus on meaningful words
    resume_words = resume_words - stop_words
    job_desc_words = job_desc_words - stop_words
    
    # Generate n-grams (bigrams and trigrams)
    resume_bigrams = set(generate_ngrams(resume_text, 2))
    job_desc_bigrams = set(generate_ngrams(job_description_text, 2))
    resume_trigrams = set(generate_ngrams(resume_text, 3))
    job_desc_trigrams = set(generate_ngrams(job_description_text, 3))

    # Matching based on n-grams and single words
    matching_keywords = resume_words.intersection(job_desc_words)
    matching_bigrams = resume_bigrams.intersection(job_desc_bigrams)
    matching_trigrams = resume_trigrams.intersection(job_desc_trigrams)

    # Combine matches
    matching_keywords.update(matching_bigrams)
    matching_keywords.update(matching_trigrams)

    return matching_keywords, len(matching_keywords), len(job_desc_words)

def check_formatting(resume_text, resume_doc=None):
    """ Check for basic ATS-friendly formatting like bullet points, clear sections """
    formatting_score = 0

    # Check for the presence of bullet points (using symbols like "-", "*", or "•")
    if '-' in resume_text or '*' in resume_text or '•' in resume_text:
        formatting_score += 1

    # Check for sections in the document (e.g., "Work Experience", "Education", etc.)
    required_sections = ['work experience', 'education', 'skills', 'certifications', 'summary']
    sections_found = sum(1 for section in required_sections if section in resume_text.lower())
    
    # Check for bold headers in .docx files
    if resume_doc:
        bold_headers_found = 0
        for para in resume_doc.paragraphs:
            if para.style.font.bold:  # Check for bold text in .docx
                for section in required_sections:
                    if section in para.text.lower():
                        bold_headers_found += 1
        formatting_score += bold_headers_found

    formatting_score += sections_found
    return formatting_score, sections_found, required_sections


def is_ats_friendly(file_path, job_description):
    """ Main function to determine if the resume is ATS-friendly """
    # Read the resume file based on the file extension
    if file_path.endswith(".docx"):
        resume_text = read_docx(file_path)
        resume_doc = Document(file_path)  # We need the document object to check bold
    elif file_path.endswith(".pdf"):
        resume_text = read_pdf(file_path)
        resume_doc = None  # PDF doesn't support direct styling checks
    else:
        raise ValueError("Unsupported file type. Please provide a .docx or .pdf file.")

    # Analyze the text
    resume_text_clean = clean_text(resume_text)
    matching_keywords, matched_keywords_count, total_keywords = check_keywords(resume_text, job_description)
    formatting_score, sections_found, required_sections = check_formatting(resume_text_clean, resume_doc)

    # Determine if it's ATS friendly based on some thresholds
    ats_friendly = True
    if matched_keywords_count / total_keywords < 0.5:
        ats_friendly = False
        print(f"Warning: Less than 50% of the job description keywords are matched ({matched_keywords_count}/{total_keywords})")
    
    if formatting_score < 3:
        ats_friendly = False
        print("Warning: Resume formatting might not be ATS-friendly. Missing bullet points or common sections.")
    
    # Display results
    print(f"\n--- ATS Resume Analysis ---")
    print(f"Matching keywords: {', '.join(matching_keywords) if matching_keywords else 'None'}")
    print(f"Formatting score: {formatting_score}/{len(required_sections)} sections found ({sections_found} of {len(required_sections)} required sections).")
    print(f"ATS Friendly: {'Yes' if ats_friendly else 'No'}")

    return ats_friendly

# Sample usage
job_description = """
    Senior Software Engineer with experience in Python, JavaScript, HTML, CSS, C++, SQL, BOotstrap, TailwindCSS.
    Expertise in creating high-performance, scalable systems.
    """
resume_file_path = r"C:\Users\nares\OneDrive\Desktop\ATS-Resume-Checker-main\naresh updated resume 007.pdf"  # Replace with the path to your resume file

is_ats_friendly(resume_file_path, job_description)